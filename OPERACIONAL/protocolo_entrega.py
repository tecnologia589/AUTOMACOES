"""
=============================================================================
  PROTOCOLO DE ENTREGA E RECEBIMENTO DE DOCUMENTOS - CORBELINO ADVOGADOS ASSOCIADOS
=============================================================================

  Gera protocolo padronizado de entrega/recebimento de documentos,
  registra publicacao no ADVBOX e (opcionalmente) sobe PDF para o Drive.

  Uso:
    # Por ID do processo
    python OPERACIONAL/protocolo_entrega.py <lawsuit_id> --tipo entrega --docs "CNH, CTPS, Comprovante"

    # Por nome do cliente (busca no ADVBOX)
    python OPERACIONAL/protocolo_entrega.py --cliente "JOAO SILVA" --tipo recebimento --docs "Sentenca, Calculo"

    # Com upload para o Drive
    python OPERACIONAL/protocolo_entrega.py <lawsuit_id> --tipo entrega --docs "CNH" --drive

    # Com registro no ADVBOX (publicacao)
    python OPERACIONAL/protocolo_entrega.py <lawsuit_id> --tipo entrega --docs "CNH" --advbox

    # Completo: gera + sobe Drive + registra ADVBOX
    python OPERACIONAL/protocolo_entrega.py <lawsuit_id> --tipo entrega --docs "CNH, CTPS" --drive --advbox

  Tipos:
    entrega      - Escritorio ENTREGA documentos ao cliente
    recebimento  - Escritorio RECEBE documentos do cliente

  Configuracao (config/.env):
    ESCRITORIO_NOME, ESCRITORIO_ADVOGADO, ESCRITORIO_OAB, ESCRITORIO_CIDADE
    ADVBOX_USER_FROM / ADVBOX_USER_RESPONSAVEL  -> campo 'from' das publicacoes
    ADVBOX_USER_OPERACIONAL                     -> guest das publicacoes
    GOOGLE_PASTA_RECLAMANTE                     -> ID da pasta RECLAMANTE no Drive
=============================================================================
"""
import sys
import os
import io
import argparse
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'INTEGRACOES'))

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), '..', 'config', '.env'))

from advbox_integration import (
    buscar_cliente, buscar_processo, obter_cliente,
    criar_publicacao, buscar_tipo_tarefa, carregar_settings
)


# ============================================================
# DADOS DO ESCRITORIO (lidos do ambiente, com defaults CORBELINO_ADVOGADOS)
# ============================================================
ESCRITORIO = {
    'nome': os.getenv('ESCRITORIO_NOME', 'CORBELINO ADVOGADOS ASSOCIADOS'),
    'advogada': os.getenv('ESCRITORIO_ADVOGADO', 'Dr. Paulo Alexandre Soares Corbelino'),
    'oab': os.getenv('ESCRITORIO_OAB', 'OAB/MT 33.267'),
    'endereco': os.getenv('ESCRITORIO_CIDADE', 'Cáceres/MT'),
}

# IDs de usuario ADVBOX - lidos do ambiente, nunca hardcoded.
USER_FROM_ID = os.getenv('ADVBOX_USER_FROM', '') or os.getenv('ADVBOX_USER_RESPONSAVEL', '')
USER_OPERACIONAL_ID = os.getenv('ADVBOX_USER_OPERACIONAL', '')

# ID da pasta RECLAMANTE no Drive - configurar no config/.env (vazio por padrao).
PASTA_RECLAMANTE_ID = os.getenv("GOOGLE_PASTA_RECLAMANTE", "")


# ============================================================
# BUSCAR DADOS DO PROCESSO E CLIENTE NO ADVBOX
# ============================================================

def obter_dados_processo(lawsuit_id=None, nome_cliente=None):
    """Busca dados do processo e cliente no ADVBOX. Retorna (processo, cliente) ou (None, None)."""
    processo = None
    cliente = None

    if lawsuit_id:
        from advbox_integration import _request
        proc_data = _request('GET', f'/lawsuits/{lawsuit_id}')
        if proc_data and proc_data.get('data'):
            processo = proc_data['data']
        elif proc_data:
            processo = proc_data
        if processo:
            customers = processo.get('customers', [])
            if customers:
                cid = customers[0].get('id') or customers[0].get('customers_id')
                if cid:
                    cli_data = obter_cliente(cid)
                    if cli_data and cli_data.get('data'):
                        cliente = cli_data['data']
                    elif cli_data:
                        cliente = cli_data

    elif nome_cliente:
        clientes = buscar_cliente(nome=nome_cliente)
        if clientes:
            cliente = clientes[0]
            cid = cliente.get('id')
            processos = buscar_processo(cliente_id=cid)
            if processos:
                processo = processos[0]

    return processo, cliente


# ============================================================
# GERAR PROTOCOLO (TEXTO)
# ============================================================

def gerar_texto_protocolo(tipo, documentos, cliente_dados, processo_dados, observacoes=None):
    """Gera texto formatado do protocolo."""
    agora = datetime.now()
    data_formatada = agora.strftime('%d/%m/%Y')
    hora_formatada = agora.strftime('%H:%M')

    nome_cliente = (cliente_dados or {}).get('name', 'NAO IDENTIFICADO')
    cpf_cliente = (cliente_dados or {}).get('identification', '')
    num_processo = (processo_dados or {}).get('process_number', 'S/N')
    pasta = (processo_dados or {}).get('folder', '')

    tipo_texto = 'ENTREGA' if tipo == 'entrega' else 'RECEBIMENTO'
    acao = 'entrega' if tipo == 'entrega' else 'recebe'
    preposicao = 'ao(a) cliente' if tipo == 'entrega' else 'do(a) cliente'

    linhas = []
    linhas.append('=' * 70)
    linhas.append(f'  {ESCRITORIO["nome"]}')
    linhas.append(f'  {ESCRITORIO["advogada"]} - {ESCRITORIO["oab"]}')
    linhas.append('=' * 70)
    linhas.append('')
    linhas.append(f'  PROTOCOLO DE {tipo_texto} DE DOCUMENTOS')
    linhas.append(f'  Data: {data_formatada}    Hora: {hora_formatada}')
    linhas.append('')
    linhas.append('-' * 70)
    linhas.append('  DADOS DO CLIENTE')
    linhas.append('-' * 70)
    linhas.append(f'  Nome:     {nome_cliente}')
    if cpf_cliente:
        linhas.append(f'  CPF:      {cpf_cliente}')
    if num_processo and num_processo != 'S/N':
        linhas.append(f'  Processo: {num_processo}')
    if pasta:
        linhas.append(f'  Pasta:    {pasta}')
    linhas.append('')
    linhas.append('-' * 70)
    linhas.append(f'  DOCUMENTOS ({tipo_texto})')
    linhas.append('-' * 70)

    for i, doc in enumerate(documentos, 1):
        linhas.append(f'  {i:>3}. {doc.strip()}')

    linhas.append('')
    linhas.append(f'  Total: {len(documentos)} documento(s)')
    linhas.append('')

    if observacoes:
        linhas.append('-' * 70)
        linhas.append('  OBSERVACOES')
        linhas.append('-' * 70)
        linhas.append(f'  {observacoes}')
        linhas.append('')

    linhas.append('-' * 70)
    linhas.append(f'  O escritorio {ESCRITORIO["nome"]} {acao} os documentos acima')
    linhas.append(f'  {preposicao} {nome_cliente}.')
    linhas.append('')
    linhas.append(f'  {ESCRITORIO["endereco"]}, {data_formatada}.')
    linhas.append('')
    linhas.append('')
    linhas.append('  ________________________________________')
    linhas.append(f'  {ESCRITORIO["advogada"]}')
    linhas.append(f'  {ESCRITORIO["oab"]}')
    linhas.append('')
    linhas.append('')
    linhas.append('  ________________________________________')
    linhas.append(f'  {nome_cliente}')
    if cpf_cliente:
        linhas.append(f'  CPF: {cpf_cliente}')
    linhas.append('')
    linhas.append('=' * 70)

    return '\n'.join(linhas)


# ============================================================
# GERAR DOCX (FORMATADO)
# ============================================================

def gerar_docx_protocolo(tipo, documentos, cliente_dados, processo_dados, observacoes=None):
    """Gera DOCX formatado do protocolo. Retorna path do arquivo gerado."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print('  AVISO: python-docx nao instalado. Instale com: pip install python-docx')
        print('  Gerando apenas versao texto...')
        return None

    agora = datetime.now()
    data_formatada = agora.strftime('%d/%m/%Y')
    hora_formatada = agora.strftime('%H:%M')

    nome_cliente = (cliente_dados or {}).get('name', 'NAO IDENTIFICADO')
    cpf_cliente = (cliente_dados or {}).get('identification', '')
    num_processo = (processo_dados or {}).get('process_number', 'S/N')
    pasta = (processo_dados or {}).get('folder', '')

    tipo_texto = 'ENTREGA' if tipo == 'entrega' else 'RECEBIMENTO'
    acao = 'entrega' if tipo == 'entrega' else 'recebe'
    preposicao = 'ao(a) cliente' if tipo == 'entrega' else 'do(a) cliente'

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Montserrat'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Cabecalho
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ESCRITORIO['nome'])
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Montserrat'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{ESCRITORIO["advogada"]} - {ESCRITORIO["oab"]}')
    run.font.size = Pt(10)
    run.font.name = 'Montserrat'

    doc.add_paragraph('')

    # Titulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'PROTOCOLO DE {tipo_texto} DE DOCUMENTOS')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Montserrat'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Data: {data_formatada}    Hora: {hora_formatada}')
    run.font.size = Pt(10)
    run.font.name = 'Montserrat'

    doc.add_paragraph('')

    # Dados do cliente
    p = doc.add_paragraph()
    run = p.add_run('DADOS DO CLIENTE')
    run.bold = True
    run.font.name = 'Montserrat'

    def add_campo(label, valor):
        if not valor:
            return
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.name = 'Montserrat'
        run = p.add_run(str(valor))
        run.font.name = 'Montserrat'

    add_campo('Nome', nome_cliente)
    add_campo('CPF', cpf_cliente)
    if num_processo and num_processo != 'S/N':
        add_campo('Processo', num_processo)
    if pasta:
        add_campo('Pasta', pasta)

    doc.add_paragraph('')

    # Lista de documentos
    p = doc.add_paragraph()
    run = p.add_run(f'DOCUMENTOS ({tipo_texto})')
    run.bold = True
    run.font.name = 'Montserrat'

    # Tabela de documentos
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['N.', 'Documento', 'Status']):
        hdr[i].text = txt
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Montserrat'
                run.font.size = Pt(10)

    for i, doc_nome in enumerate(documentos, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = doc_nome.strip()
        status = 'Entregue' if tipo == 'entrega' else 'Recebido'
        row[2].text = status
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run(f'Total: {len(documentos)} documento(s)')
    run.font.name = 'Montserrat'

    # Observacoes
    if observacoes:
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run('OBSERVACOES')
        run.bold = True
        run.font.name = 'Montserrat'
        p = doc.add_paragraph()
        run = p.add_run(observacoes)
        run.font.name = 'Montserrat'

    doc.add_paragraph('')

    # Declaracao
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        f'O escritorio {ESCRITORIO["nome"]} {acao} os documentos acima listados '
        f'{preposicao} {nome_cliente}, conforme protocolo.'
    )
    run.font.name = 'Montserrat'

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'{ESCRITORIO["endereco"]}, {data_formatada}.')
    run.font.name = 'Montserrat'

    # Assinaturas
    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('________________________________________')
    run.font.name = 'Montserrat'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{ESCRITORIO["advogada"]}')
    run.bold = True
    run.font.name = 'Montserrat'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ESCRITORIO['oab'])
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('________________________________________')
    run.font.name = 'Montserrat'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome_cliente)
    run.bold = True
    run.font.name = 'Montserrat'
    if cpf_cliente:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'CPF: {cpf_cliente}')
        run.font.name = 'Montserrat'
        run.font.size = Pt(10)

    # Salvar
    nome_arq = f'{nome_cliente} - Protocolo {tipo_texto} - {agora.strftime("%d.%m.%Y")}.docx'
    nome_arq = nome_arq.replace('/', '-').replace('\\', '-')
    caminho = os.path.join(os.path.dirname(__file__), '..', nome_arq)
    doc.save(caminho)
    return caminho


# ============================================================
# UPLOAD PARA O DRIVE
# ============================================================

def subir_protocolo_drive(caminho_docx, nome_cliente):
    """Sobe protocolo DOCX para a pasta do cliente no Drive (ATOS INTERNOS)."""
    from INTEGRACOES.google_integration import autenticar_google

    if not PASTA_RECLAMANTE_ID:
        print('  AVISO: GOOGLE_PASTA_RECLAMANTE nao configurado no config/.env.')
        print(f'  Protocolo salvo apenas localmente: {caminho_docx}')
        return None

    drive_service, _ = autenticar_google()

    # Buscar pasta do cliente em RECLAMANTE
    query = (
        f"'{PASTA_RECLAMANTE_ID}' in parents "
        f"and mimeType='application/vnd.google-apps.folder' "
        f"and trashed=false"
    )
    results = drive_service.files().list(
        q=query, fields='files(id, name)', pageSize=500
    ).execute()

    pasta_cliente_id = None
    for pasta in results.get('files', []):
        if nome_cliente.upper() in pasta['name'].upper():
            pasta_cliente_id = pasta['id']
            print(f'  Pasta encontrada: {pasta["name"]}')
            break

    if not pasta_cliente_id:
        print(f'  AVISO: Pasta do cliente "{nome_cliente}" nao encontrada no Drive.')
        print(f'  Protocolo salvo apenas localmente: {caminho_docx}')
        return None

    # Buscar subpasta ATOS INTERNOS
    query_sub = (
        f"'{pasta_cliente_id}' in parents "
        f"and mimeType='application/vnd.google-apps.folder' "
        f"and name='ATOS INTERNOS' and trashed=false"
    )
    sub_results = drive_service.files().list(q=query_sub, fields='files(id)').execute()
    subpastas = sub_results.get('files', [])
    pasta_destino = subpastas[0]['id'] if subpastas else pasta_cliente_id

    # Upload como Google Docs (converte automaticamente)
    from googleapiclient.http import MediaFileUpload
    nome_arquivo = os.path.basename(caminho_docx)
    media = MediaFileUpload(caminho_docx, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file_metadata = {
        'name': nome_arquivo.replace('.docx', ''),
        'parents': [pasta_destino],
        'mimeType': 'application/vnd.google-apps.document',
    }
    uploaded = drive_service.files().create(
        body=file_metadata, media_body=media, fields='id, webViewLink'
    ).execute()

    link = uploaded.get('webViewLink', '')
    print(f'  Protocolo no Drive: {link}')
    return link


# ============================================================
# REGISTRAR NO ADVBOX (PUBLICACAO)
# ============================================================

def registrar_protocolo_advbox(lawsuit_id, tipo, documentos, link_drive=None):
    """Registra protocolo como publicacao no ADVBOX."""
    tipo_texto = 'ENTREGA' if tipo == 'entrega' else 'RECEBIMENTO'
    data_hoje = datetime.now().strftime('%d/%m/%Y')

    docs_lista = '\n'.join(f'  - {d.strip()}' for d in documentos)
    comentario = (
        f'PROTOCOLO DE {tipo_texto} DE DOCUMENTOS - {data_hoje}\n\n'
        f'Documentos:\n{docs_lista}\n\n'
        f'Total: {len(documentos)} documento(s)'
    )
    if link_drive:
        comentario += f'\n\nLink: {link_drive}'

    # Buscar tipo de tarefa ACOMPANHAMENTO
    task_id = buscar_tipo_tarefa('ACOMPANHAMENTO')
    if not task_id:
        task_id = buscar_tipo_tarefa('ANDAMENTO')
    if not task_id:
        # Pega o primeiro disponivel
        settings = carregar_settings()
        tasks = settings.get('tasks', [])
        task_id = tasks[0]['id'] if tasks else None

    if not task_id:
        print('  ERRO: Nenhum tipo de tarefa encontrado no ADVBOX.')
        return None

    if not USER_FROM_ID:
        print('  ERRO: ADVBOX_USER_FROM/ADVBOX_USER_RESPONSAVEL nao configurado no config/.env')
        return None

    guest_ids = [int(USER_OPERACIONAL_ID)] if USER_OPERACIONAL_ID else []

    resultado = criar_publicacao(
        lawsuit_id=lawsuit_id,
        task_id=task_id,
        guest_ids=guest_ids,
        comments=comentario,
        from_id=int(USER_FROM_ID),
    )

    if resultado and resultado.get('success'):
        print(f'  Publicacao registrada no ADVBOX! ID: {resultado.get("posts_id")}')
    elif resultado:
        print(f'  Publicacao criada. Resposta: {resultado}')
    return resultado


# ============================================================
# COMANDO PRINCIPAL
# ============================================================

def executar_protocolo(args):
    """Executa fluxo completo de protocolo."""
    print('\n' + '=' * 70)
    print('  PROTOCOLO DE ENTREGA/RECEBIMENTO - CORBELINO ADVOGADOS ASSOCIADOS')
    print('=' * 70)

    # Parsear documentos
    documentos = [d.strip() for d in args.docs.split(',') if d.strip()]
    if not documentos:
        print('  ERRO: Informe ao menos um documento com --docs "Doc1, Doc2"')
        sys.exit(1)

    tipo = args.tipo.lower()
    if tipo not in ('entrega', 'recebimento'):
        print('  ERRO: --tipo deve ser "entrega" ou "recebimento"')
        sys.exit(1)

    # Buscar dados no ADVBOX
    print(f'\n  Buscando dados no ADVBOX...')
    lawsuit_id = args.processo if hasattr(args, 'processo') and args.processo else None
    processo, cliente = obter_dados_processo(
        lawsuit_id=lawsuit_id,
        nome_cliente=args.cliente if hasattr(args, 'cliente') and args.cliente else None
    )

    nome_cliente = (cliente or {}).get('name', args.cliente or 'CLIENTE')
    tipo_texto = 'ENTREGA' if tipo == 'entrega' else 'RECEBIMENTO'

    if processo:
        print(f'  Processo: {processo.get("process_number", "S/N")}')
        if not lawsuit_id:
            lawsuit_id = processo.get('id')
    if cliente:
        print(f'  Cliente: {cliente.get("name", "?")}')

    print(f'  Tipo: {tipo_texto}')
    print(f'  Documentos: {len(documentos)}')
    for d in documentos:
        print(f'    - {d}')

    # Gerar texto
    texto = gerar_texto_protocolo(tipo, documentos, cliente, processo, args.obs)
    print(f'\n{texto}')

    # Gerar DOCX
    caminho_docx = gerar_docx_protocolo(tipo, documentos, cliente, processo, args.obs)
    if caminho_docx:
        print(f'\n  DOCX gerado: {caminho_docx}')

    # Upload Drive
    link_drive = None
    if args.drive and caminho_docx:
        print(f'\n  Subindo para o Google Drive...')
        link_drive = subir_protocolo_drive(caminho_docx, nome_cliente)

    # Registrar no ADVBOX
    if args.advbox and lawsuit_id:
        print(f'\n  Registrando publicacao no ADVBOX...')
        registrar_protocolo_advbox(lawsuit_id, tipo, documentos, link_drive)
    elif args.advbox and not lawsuit_id:
        print(f'\n  AVISO: Sem lawsuit_id, nao foi possivel registrar no ADVBOX.')

    print('\n' + '=' * 70)
    print('  PROTOCOLO CONCLUIDO!')
    print('=' * 70)

    return caminho_docx


# ============================================================
# MAIN
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='Protocolo de Entrega/Recebimento de Documentos - Corbelino Advogados Associados'
    )
    parser.add_argument('processo', nargs='?', help='ID do processo no ADVBOX')
    parser.add_argument('--cliente', '-c', help='Nome do cliente (busca no ADVBOX)')
    parser.add_argument('--tipo', '-t', required=True, choices=['entrega', 'recebimento'],
                        help='Tipo: entrega ou recebimento')
    parser.add_argument('--docs', '-d', required=True,
                        help='Lista de documentos separados por virgula')
    parser.add_argument('--obs', '-o', help='Observacoes adicionais')
    parser.add_argument('--drive', action='store_true', help='Subir protocolo para o Google Drive')
    parser.add_argument('--advbox', action='store_true', help='Registrar publicacao no ADVBOX')

    args = parser.parse_args()

    if not args.processo and not args.cliente:
        print('ERRO: Informe o ID do processo ou --cliente "Nome"')
        sys.exit(1)

    executar_protocolo(args)


if __name__ == '__main__':
    main()
