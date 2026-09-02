"""
Motor de templates DOCX para gerar pecas no timbrado do escritorio.

Abre o template (com cabecalho/rodape/logo do escritorio), apaga o conteudo
do corpo (preservando section, headers e footers) e insere o texto gerado
pela IA com formatacao padrao do escritorio:
  - Fonte Montserrat 11pt
  - Justificado
  - Espacamento 1.5
  - Recuo de primeira linha 7pt (ajustavel)
  - Negrito via **texto**
  - Italico via _texto_
  - Titulos: linha INTEIRAMENTE EM MAIUSCULAS = titulo (negrito, sem recuo)

O timbrado e fornecido pelo escritorio em config/timbrado_modelo.docx.
"""
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import TIMBRADO_PATH

log = logging.getLogger('agente_op.template')

TEMPLATE_PATH = TIMBRADO_PATH

FONTE = 'Montserrat'
TAM_CORPO = Pt(11)
RECUO_PRIMEIRA = Cm(0.7)


def _limpar_corpo(doc):
    """Remove todos os paragrafos do body preservando section, header, footer."""
    body = doc.element.body
    # remove todos os paragrafos e tabelas (mantem o sectPr - section properties)
    for el in list(body):
        tag = el.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            body.remove(el)


def _aplicar_estilo_run(run, bold=False, italic=False):
    run.font.name = FONTE
    run.font.size = TAM_CORPO
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def _add_paragrafo(doc, texto, *, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   negrito_total=False, italic_total=False, recuo=True,
                   espaco_depois=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alinhamento
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = espaco_depois
    if recuo:
        pf.first_line_indent = RECUO_PRIMEIRA

    # Parse inline: **bold** e _italic_
    # tokens alternam: texto normal / **bold** / _italic_
    pos = 0
    pattern = re.compile(r'(\*\*([^*]+)\*\*|_([^_\n]+)_)')
    for m in pattern.finditer(texto):
        if m.start() > pos:
            r = p.add_run(texto[pos:m.start()])
            _aplicar_estilo_run(r, bold=negrito_total, italic=italic_total)
        if m.group(2):  # **bold**
            r = p.add_run(m.group(2))
            _aplicar_estilo_run(r, bold=True, italic=italic_total)
        elif m.group(3):  # _italic_
            r = p.add_run(m.group(3))
            _aplicar_estilo_run(r, bold=negrito_total, italic=True)
        pos = m.end()
    if pos < len(texto):
        r = p.add_run(texto[pos:])
        _aplicar_estilo_run(r, bold=negrito_total, italic=italic_total)
    return p


# Nomes de pecas que viram TITULO PRINCIPAL CENTRALIZADO (apenas o tipo da peca)
NOMES_PECA_CENTRALIZADO = {
    'RECLAMACAO TRABALHISTA', 'RECLAMAÇÃO TRABALHISTA',
    'CONTESTACAO', 'CONTESTAÇÃO',
    'REPLICA', 'RÉPLICA',
    'RAZOES FINAIS', 'RAZÕES FINAIS',
    'MEMORIAIS', 'MEMORIAIS ESCRITOS',
    'RECURSO ORDINARIO', 'RECURSO ORDINÁRIO',
    'RECURSO ESPECIAL', 'RECURSO EXTRAORDINARIO', 'RECURSO EXTRAORDINÁRIO',
    'CONTRARRAZOES', 'CONTRARRAZÕES',
    'CONTRARRAZOES AO RECURSO ORDINARIO', 'CONTRARRAZÕES AO RECURSO ORDINÁRIO',
    'CONTRARRAZOES AO RECURSO ESPECIAL', 'CONTRARRAZÕES AO RECURSO ESPECIAL',
    'CONTRARRAZOES AO RECURSO EXTRAORDINARIO', 'CONTRARRAZÕES AO RECURSO EXTRAORDINÁRIO',
    'CONTRARRAZOES AO RECURSO DE REVISTA', 'CONTRARRAZÕES AO RECURSO DE REVISTA',
    'RAZOES FINAIS', 'RAZÕES FINAIS', 'MEMORIAIS', 'RECURSO DE REVISTA',
    'EMBARGOS DE DECLARACAO', 'EMBARGOS DE DECLARAÇÃO',
    'AGRAVO DE INSTRUMENTO', 'AGRAVO INTERNO',
    'PETICAO INICIAL', 'PETIÇÃO INICIAL',
    'MANIFESTACAO', 'MANIFESTAÇÃO',
    'PARECER', 'PARECER JURIDICO', 'PARECER JURÍDICO',
    'ANALISE DE SENTENCA', 'ANÁLISE DE SENTENÇA',
}


def _normalizar_acento(s):
    """Pra comparacao com NOMES_PECA_CENTRALIZADO."""
    return (s.upper().replace('Ç', 'C').replace('Á', 'A').replace('Ã', 'A')
            .replace('É', 'E').replace('Ê', 'E').replace('Í', 'I').replace('Ó', 'O')
            .replace('Õ', 'O').replace('Ô', 'O').replace('Ú', 'U'))


def _eh_titulo_principal(linha):
    """True se for o NOME da peca (ex: RECLAMACAO TRABALHISTA) - vira centralizado."""
    s = linha.strip()
    if not s:
        return False
    # tira pontuacao do final
    s_limpo = re.sub(r'[\.\:\-—\s]+$', '', s)
    norm = _normalizar_acento(s_limpo)
    return norm in NOMES_PECA_CENTRALIZADO


# Padrao de titulo de SECAO (justificado + negrito, NAO centralizado):
#   "I. DOS FATOS", "II. DO DIREITO", "I.I. DA JUSTICA GRATUITA",
#   "1. DOS FATOS", "1.1 DA NULIDADE", "OMISSAO 1 — DA NULIDADE"
# NAO inclui letra-paren A) B) (esses sao subtitulos)
_RE_TITULO_SECAO_NUM = re.compile(
    r'^(?:[IVXLCDM]+(?:\.[IVXLCDM\d]+)*\.?|\d+(?:\.\d+)*\.?)\s+\S',
)
# Subtitulo: A) Das Horas Extras, B) Da Supressao, 1) Do Pedido
_RE_SUBTITULO = re.compile(r'^[A-Za-z0-9]\)\s+[A-Za-zÀ-ÿ]')


def _eh_titulo_secao(linha):
    """True se for titulo de secao (I., II., I.I., 1., 1.1, OMISSAO N —) - JUSTIFICADO + negrito.
    Aceita tambem o padrao 'PALAVRA NUM —' (ex: OMISSAO 1 —)."""
    s = linha.strip()
    if not s or len(s) < 5 or len(s) > 250:
        return False
    # padrao 1: numeracao romana ou arabe ("I. DOS FATOS", "1.1 DA NULIDADE")
    if _RE_TITULO_SECAO_NUM.match(s):
        # tem que ter texto, nao so numeracao
        resto = re.sub(r'^(?:[IVXLCDM]+(?:\.[IVXLCDM\d]+)*\.?|\d+(?:\.\d+)*\.?)\s+', '', s)
        return len(resto) >= 3
    # padrao 2: "PALAVRA NUM —" (OMISSAO 1 —, TESE 2 —, PEDIDO 3 —)
    # primeira palavra em maiusculo + numero + travessao
    if re.match(r'^[A-ZÀ-Þ]{3,}\s+\d+\s*[—\-:]\s+\S', s):
        return True
    return False


def _eh_subtitulo(linha):
    """True se for subtitulo do tipo 'A) Das ...', 'B) Da ...'."""
    s = linha.strip()
    if not s or len(s) < 5 or len(s) > 150:
        return False
    return bool(_RE_SUBTITULO.match(s))


def _eh_titulo(linha):
    """Compatibilidade: titulo = principal OU secao OU subtitulo."""
    return _eh_titulo_principal(linha) or _eh_titulo_secao(linha) or _eh_subtitulo(linha)


# Padroes que indicam citacao em bloco (jurisprudencia, sumula, doutrina, lei transcrita)
# Aceita: "Sumula NNN", "OJ NNN", "Tese NNN", "Tema NNN", "Art. NNN ... —", "Art. NNN, X, da Lei ..."
_RE_CITACAO_INICIO = re.compile(
    r'^(?:["“«]|s[uú]mula\s+\d+|oj\s+\d+|tese\s+\d+|tema\s+\d+'
    r'|art(?:igo)?\.?\s+\d+(?:[,\.\-\s][^—\-:]*)?[—\-:])',
    re.IGNORECASE,
)


def _eh_citacao(linha):
    """Detecta citacao em bloco:
    - Comeca com aspas (", ", «)
    - OU comeca com 'Sumula NNN', 'OJ NNN', 'Artigo NNN —', 'Tese NNN', 'Tema NNN'
      (texto razoavel >30 chars, evita falso positivo de referencia inline)
    """
    s = linha.strip()
    if not s:
        return False
    if s.startswith(('"', '"', '“', '«')):
        return True
    if len(s) >= 30 and _RE_CITACAO_INICIO.match(s):
        return True
    return False


def _adicionar_imagens_anexo(doc, imagens, drive_service):
    """
    Adiciona seção de ANEXOS ao final do docx com as imagens da pasta do cliente.
    imagens: lista de {id, name, mime}
    drive_service: client Google Drive autenticado
    """
    import io as _io
    from docx.enum.text import WD_BREAK
    from googleapiclient.http import MediaIoBaseDownload
    if not imagens:
        return

    # Quebra de pagina antes dos anexos
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # Titulo da secao de anexos
    _add_paragrafo(doc, 'ANEXOS - PROVAS DOCUMENTAIS (IMAGENS)',
                   alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                   negrito_total=True, recuo=False, espaco_depois=Pt(14))

    for i, img in enumerate(imagens, 1):
        try:
            # Download
            buf = _io.BytesIO()
            req = drive_service.files().get_media(fileId=img['id'], supportsAllDrives=True)
            dl = MediaIoBaseDownload(buf, req)
            done = False
            while not done:
                _, done = dl.next_chunk()
            buf.seek(0)

            # Legenda
            _add_paragrafo(doc, f'**ANEXO {i}** — {img["name"]}',
                           alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                           recuo=False, espaco_depois=Pt(4))

            # Imagem (max 15cm largura)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            try:
                run.add_picture(buf, width=Cm(15))
            except Exception as e:
                log.warning(f'falha ao inserir imagem {img["name"]}: {e}')
                continue
            log.info(f'  anexo {i} inserido: {img["name"]}')

            # Espaco apos
            doc.add_paragraph()
        except Exception as e:
            log.warning(f'falha ao baixar/inserir imagem {img["name"]}: {e}')


def gerar_peca_no_template(texto_peca: str, output_path: str | Path,
                           imagens=None, drive_service=None) -> Path:
    """
    Gera um DOCX no timbrado do escritorio com o conteudo da peca.

    Convencoes do texto_peca:
      - Linhas em branco separam paragrafos
      - Linha INTEIRAMENTE EM MAIUSCULAS vira titulo (negrito, centralizado)
      - Linhas iniciando com aspas viram citacoes (recuo 4cm, italico)
      - **texto** -> negrito inline
      - _texto_ -> italico inline
    """
    if not Path(TEMPLATE_PATH).exists():
        raise FileNotFoundError(
            f'Timbrado nao encontrado: {TEMPLATE_PATH}. '
            f'O escritorio deve fornecer o timbrado DELE em config/timbrado_modelo.docx.'
        )

    doc = Document(str(TEMPLATE_PATH))
    _limpar_corpo(doc)

    paragrafos = [p.strip() for p in texto_peca.split('\n')]
    # consolida quebras: separa por linhas em branco
    blocos = []
    atual = []
    for linha in paragrafos:
        if linha:
            atual.append(linha)
        else:
            if atual:
                blocos.append('\n'.join(atual))
                atual = []
    if atual:
        blocos.append('\n'.join(atual))

    for bloco in blocos:
        linha_unica = '\n' not in bloco
        # 1) NOME DA PECA (RECLAMACAO TRABALHISTA, CONTESTACAO, etc) -> centralizado
        if linha_unica and _eh_titulo_principal(bloco):
            _add_paragrafo(doc, bloco,
                           alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                           negrito_total=True, recuo=False,
                           espaco_depois=Pt(14))
        # 2) TITULO DE SECAO (I., II., I.I., 1., 1.1) -> justificado + negrito
        elif linha_unica and _eh_titulo_secao(bloco):
            _add_paragrafo(doc, bloco,
                           alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           negrito_total=True, recuo=False,
                           espaco_depois=Pt(12))
        # 3) SUBTITULO (A) Das Horas, B) Da Supressao) -> justificado + negrito
        elif linha_unica and _eh_subtitulo(bloco):
            _add_paragrafo(doc, bloco,
                           alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           negrito_total=True, recuo=False,
                           espaco_depois=Pt(8))
        # 4) CITACAO (sumula, jurisprudencia, etc) -> recuo 4cm + italico
        elif _eh_citacao(bloco):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.left_indent = Cm(4)
            pf.space_after = Pt(6)
            r = p.add_run(bloco)
            _aplicar_estilo_run(r, italic=True)
        # 5) PARAGRAFO COMUM
        else:
            _add_paragrafo(doc, bloco)

    # Adiciona imagens como anexos no final (se houver)
    if imagens and drive_service:
        _adicionar_imagens_anexo(doc, imagens, drive_service)

    out = Path(output_path)
    doc.save(str(out))
    log.info(f'Peca salva no timbrado: {out} (imagens={len(imagens) if imagens else 0})')
    return out
