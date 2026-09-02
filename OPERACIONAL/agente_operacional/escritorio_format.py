"""
Motor de formatacao FIEL ao padrao de pecas do escritorio.

Especificacoes (padrao do escritorio):
  Pagina A4
  Margens: top 3cm | bottom 1.6cm | left 3cm | right 3cm
  Fonte: Montserrat
  Corpo: 11 pt
  Citacoes: 10 pt italico
  Espacamento entre linhas: 1.5
  Recuo 1a linha do corpo: 7 cm (literal)
  Recuo esquerdo das citacoes: 4 cm
  space_before e space_after: None (sem espaco explicito entre paragrafos)
  Alinhamento padrao: justificado
  Endereçamento: bold, justificado, sem recuo
  Nome da peca (RECLAMATORIA TRABALHISTA): centralizado, bold
  Titulos de secao (I - DA ...): justificado, bold, sem recuo, travessao longo
  Citacoes: justificado, recuo esquerdo 4cm, 10pt, italico

Mantem o timbrado (cabecalho/logo/rodape) intacto - apenas limpa o body.
O timbrado e fornecido pelo escritorio em config/timbrado_modelo.docx.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Timbrado fornecido pelo escritorio (placeholder: config/timbrado_modelo.docx)
_ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE_PATH = _ROOT / 'config' / 'timbrado_modelo.docx'
REFERENCIAS_DIR = Path(__file__).parent / 'REFERENCIAS'

FONTE = 'Montserrat'
TAM_CORPO = Pt(11)
TAM_CITACAO = Pt(10)
RECUO_PRIMEIRA = Cm(7.0)        # 7cm literais (padrao do escritorio)
RECUO_CITACAO_ESQ = Cm(4.0)
LINE_SPACING = 1.5
MARG_TOP = Cm(3.0)
MARG_BOT = Cm(1.6)
MARG_LEFT = Cm(3.0)
MARG_RIGHT = Cm(3.0)


# ============================================================
# DETECTORES DE TIPO DE BLOCO
# ============================================================

NOMES_PECA_CENTRALIZADO = {
    'RECLAMACAO TRABALHISTA', 'RECLAMAÇÃO TRABALHISTA',
    'RECLAMATORIA TRABALHISTA', 'RECLAMATÓRIA TRABALHISTA',
    'CONTESTACAO', 'CONTESTAÇÃO',
    'REPLICA', 'RÉPLICA',
    'RAZOES FINAIS', 'RAZÕES FINAIS',
    'MEMORIAIS', 'MEMORIAIS ESCRITOS',
    'RECURSO ORDINARIO', 'RECURSO ORDINÁRIO',
    'RECURSO DE REVISTA',
    'CONTRARRAZOES', 'CONTRARRAZÕES',
    'EMBARGOS DE DECLARACAO', 'EMBARGOS DE DECLARAÇÃO',
    'AGRAVO DE INSTRUMENTO', 'AGRAVO INTERNO',
    'PETICAO INICIAL', 'PETIÇÃO INICIAL',
    'MANIFESTACAO', 'MANIFESTAÇÃO',
    'PARECER', 'PARECER JURIDICO', 'PARECER JURÍDICO',
}


def _normalizar_acento(s):
    return (s.upper().replace('Ç', 'C').replace('Á', 'A').replace('Ã', 'A')
            .replace('É', 'E').replace('Ê', 'E').replace('Í', 'I').replace('Ó', 'O')
            .replace('Õ', 'O').replace('Ô', 'O').replace('Ú', 'U'))


def _eh_nome_peca(linha):
    s = linha.strip()
    s_limpo = re.sub(r'[\.\:\-—\s]+$', '', s)
    return _normalizar_acento(s_limpo) in NOMES_PECA_CENTRALIZADO


_RE_TITULO_SECAO_NUM = re.compile(
    r'^(?:[IVXLCDM]+(?:\.[IVXLCDM\d]+)*\.?|\d+(?:\.\d+)*\.?)\s+\S',
)
_RE_SUBTITULO = re.compile(r'^[A-Za-z0-9]\)\s+[A-Za-zÀ-ÿ]')


def _eh_titulo_secao(linha):
    s = linha.strip()
    if not s or len(s) < 5 or len(s) > 250:
        return False
    if _RE_TITULO_SECAO_NUM.match(s):
        resto = re.sub(r'^(?:[IVXLCDM]+(?:\.[IVXLCDM\d]+)*\.?|\d+(?:\.\d+)*\.?)\s+', '', s)
        return len(resto) >= 3
    return False


def _eh_subtitulo(linha):
    s = linha.strip()
    if not s or len(s) < 5 or len(s) > 200:
        return False
    return bool(_RE_SUBTITULO.match(s))


def _eh_enderecamento(linha):
    s = linha.strip()
    return s.upper().startswith(('EXCELENTISSIMO', 'EXCELENTÍSSIMO', 'EXMO', 'EXMA'))


def _eh_pedido_final(linha):
    # padrao "a)", "b)", "aa)", "bb)" no inicio
    return bool(re.match(r'^[a-z]{1,2}\)\s+', linha.strip()))


_RE_CITACAO_INICIO = re.compile(
    r'^(?:["“«]|s[uú]mula\s+\d+|oj\s+\d+|tese\s+\d+|tema\s+\d+'
    r'|art(?:igo)?\.?\s+\d+(?:[,\.\-\s][^—\-:]*)?[—\-:])',
    re.IGNORECASE,
)


def _eh_citacao(linha):
    s = linha.strip()
    if not s:
        return False
    if s.startswith(('"', '"', '“', '«')):
        return True
    if len(s) >= 30 and _RE_CITACAO_INICIO.match(s):
        return True
    return False


# ============================================================
# RENDERIZADORES
# ============================================================

def _ajustar_margens(doc):
    for s in doc.sections:
        s.top_margin = MARG_TOP
        s.bottom_margin = MARG_BOT
        s.left_margin = MARG_LEFT
        s.right_margin = MARG_RIGHT


def _limpar_corpo(doc):
    body = doc.element.body
    for el in list(body):
        tag = el.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            body.remove(el)


def _estilo_run(run, *, size=TAM_CORPO, bold=False, italic=False):
    run.font.name = FONTE
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def _padrao_paragrafo(p, *, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=True, left_indent=None):
    p.alignment = alinhamento
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = None
    pf.space_after = None
    if recuo_primeira:
        pf.first_line_indent = RECUO_PRIMEIRA
    if left_indent is not None:
        pf.left_indent = left_indent


def _add_texto_inline(p, texto, *, size=TAM_CORPO, negrito_total=False, italic_total=False):
    """Adiciona texto a um paragrafo respeitando markup **bold** e _italic_."""
    pos = 0
    pattern = re.compile(r'(\*\*([^*]+)\*\*|_([^_\n]+)_)')
    for m in pattern.finditer(texto):
        if m.start() > pos:
            r = p.add_run(texto[pos:m.start()])
            _estilo_run(r, size=size, bold=negrito_total, italic=italic_total)
        if m.group(2):
            r = p.add_run(m.group(2))
            _estilo_run(r, size=size, bold=True, italic=italic_total)
        elif m.group(3):
            r = p.add_run(m.group(3))
            _estilo_run(r, size=size, bold=negrito_total, italic=True)
        pos = m.end()
    if pos < len(texto):
        r = p.add_run(texto[pos:])
        _estilo_run(r, size=size, bold=negrito_total, italic=italic_total)


def _add_paragrafo_corpo(doc, texto):
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=True)
    _add_texto_inline(p, texto)
    return p


def _add_paragrafo_titulo_secao(doc, texto):
    """Titulo de secao (I -, I.I -, etc): justificado, bold, sem recuo, travessao longo."""
    # Normaliza hifen para travessao longo na numeracao "I - "
    texto = re.sub(r'^([IVX\d\.]+)\s*[-]\s+', r'\1 – ', texto, count=1)
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=False)
    _add_texto_inline(p, texto, negrito_total=True)
    return p


def _add_paragrafo_subtitulo(doc, texto):
    """Subtitulo a), b), c): justificado, bold, sem recuo (mesmo padrao secao)."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=False)
    _add_texto_inline(p, texto, negrito_total=True)
    return p


def _add_paragrafo_nome_peca(doc, texto):
    """Nome da peca (RECLAMATORIA TRABALHISTA): centralizado, bold."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo_primeira=False)
    _add_texto_inline(p, texto, negrito_total=True)
    return p


def _add_paragrafo_enderecamento(doc, texto):
    """EXCELENTISSIMO ... : justificado, bold, sem recuo."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=False)
    _add_texto_inline(p, texto, negrito_total=True)
    return p


def _add_paragrafo_pedido(doc, texto):
    """Pedidos finais a), b), c): justificado, recuo 7cm."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo_primeira=True)
    _add_texto_inline(p, texto)
    return p


def _add_paragrafo_citacao(doc, texto):
    """Citacao em bloco: justificado, recuo esquerdo 4cm, 10pt, italico."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     recuo_primeira=False, left_indent=RECUO_CITACAO_ESQ)
    # Tudo em 10pt italico
    r = p.add_run(texto)
    _estilo_run(r, size=TAM_CITACAO, italic=True)
    return p


# ============================================================
# RENDERIZADOR PRINCIPAL
# ============================================================

def gerar_peca_escritorio(texto_peca: str, output_path):
    """Gera DOCX no timbrado seguindo a formatacao padrao do escritorio."""
    if not Path(TEMPLATE_PATH).exists():
        raise FileNotFoundError(
            f'Timbrado nao encontrado: {TEMPLATE_PATH}. '
            f'O escritorio deve fornecer o timbrado DELE em config/timbrado_modelo.docx.'
        )

    doc = Document(str(TEMPLATE_PATH))
    _ajustar_margens(doc)
    _limpar_corpo(doc)

    # Quebra em blocos (separados por linhas em branco)
    linhas = [l.rstrip() for l in texto_peca.split('\n')]
    blocos = []
    atual = []
    for l in linhas:
        if l.strip():
            atual.append(l)
        else:
            if atual:
                blocos.append('\n'.join(atual))
                atual = []
    if atual:
        blocos.append('\n'.join(atual))

    for bloco in blocos:
        unica = '\n' not in bloco
        if unica and _eh_enderecamento(bloco):
            _add_paragrafo_enderecamento(doc, bloco)
        elif unica and _eh_nome_peca(bloco):
            _add_paragrafo_nome_peca(doc, bloco)
        elif unica and _eh_titulo_secao(bloco):
            _add_paragrafo_titulo_secao(doc, bloco)
        elif unica and _eh_subtitulo(bloco):
            _add_paragrafo_subtitulo(doc, bloco)
        elif unica and _eh_pedido_final(bloco):
            _add_paragrafo_pedido(doc, bloco)
        elif _eh_citacao(bloco):
            _add_paragrafo_citacao(doc, bloco)
        else:
            _add_paragrafo_corpo(doc, bloco)

    out = Path(output_path)
    doc.save(str(out))
    return out
